from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report_output"
ASSETS = OUT / "assets"
DOCX = OUT / "ArogyaSahaya_Project_Report.docx"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)


SCREENSHOTS = [
    ("Vital Log screen", Path(r"C:\Users\navee\OneDrive\Desktop\WhatsApp Image 2026-05-15 at 10.41.34 AM.jpeg")),
    ("Medicines screen", Path(r"C:\Users\navee\OneDrive\Desktop\WhatsApp Image 2026-05-15 at 10.43.33 AM.jpeg")),
    ("Emergency SOS screen", Path(r"C:\Users\navee\OneDrive\Desktop\WhatsApp Image 2026-05-15 at 10.41.33 AM.jpeg")),
    ("Home dashboard screen", Path(r"C:\Users\navee\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\B52D385EB36085D59EFCE89F7CF80A39B3DB77E5\transfers\2026-20\WhatsApp Image 2026-05-15 at 10.43.33 AM (2).jpeg")),
    ("Arogya Mitra AI help screen", Path(r"C:\Users\navee\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\B52D385EB36085D59EFCE89F7CF80A39B3DB77E5\transfers\2026-20\WhatsApp Image 2026-05-15 at 10.41.34 AM.jpeg")),
]


GREEN = "1B6B2A"
BLUE = "1769C2"
ORANGE = "F58220"
RED = "D93030"
LIGHT_GREEN = "EAF6EE"
LIGHT_BLUE = "EAF2FB"
LIGHT_ORANGE = "FFF3E5"
GRAY = "F4F6F8"
DARK = "1F2937"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9DEE5"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.05

    for name, size, color in [
        ("Title", 25, GREEN),
        ("Heading 1", 16, GREEN),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 11, DARK),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "ArogyaSahaya V2 - Project Report"
        header.style = doc.styles["Normal"]
        header.runs[0].font.size = Pt(8.5)
        header.runs[0].font.color.rgb = RGBColor.from_string("6B7280")
        footer = section.footer.paragraphs[0]
        add_page_number(footer)
        for run in footer.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string("6B7280")


def paragraph(doc, text="", bold=False, color=None, size=None, align=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3.5)
    run = p.add_run("- " + text)
    run.font.size = Pt(9.8)
    return p


def add_banner(doc, title, subtitle=None, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "FFFFFF")
    set_cell_margins(cell, 170, 190, 160, 190)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14.5)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string("4B5563")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_key_value_table(doc, rows, widths=(3.8, 11.7), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if header:
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        c = row.cells[0]
        set_cell_shading(c, GREEN)
        set_cell_border(c, GREEN)
        set_cell_margins(c, 120, 150, 120, 150)
        r = c.paragraphs[0].add_run(header)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for label, value in rows:
        cells = table.add_row().cells
        for i, width in enumerate(widths):
            cells[i].width = Cm(width)
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_BLUE)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].text = value
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_traceability(doc, title, compact=False):
    add_tight_heading(doc, "Traceability and evaluation coverage", color=BLUE)
    rows = [
        ("Patient value", f"This section explains how {title.lower()} contributes to safer and simpler day-to-day care."),
        ("Implementation link", "The described behavior is mapped to Android screens, Room entities, permissions, services, or repository flow as applicable."),
        ("Verification focus", "Manual validation checks visible output, data persistence, navigation behavior, and edge cases for empty or missing records."),
    ]
    if not compact:
        rows.extend([
            ("Risk control", "Health guidance remains supportive, emergency actions remain prominent, and clinical decisions are left to qualified professionals."),
            ("Documentation value", "The section gives evaluators enough context to understand both the feature purpose and its technical basis."),
        ])
    add_matrix_table(doc, ["Coverage Area", "Documented Detail"], rows, widths=[4.0, 11.2], accent=BLUE)


def add_matrix_table(doc, headers, rows, widths=None, accent=GREEN):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        if widths:
            hdr[i].width = Cm(widths[i])
        set_cell_shading(hdr[i], accent)
        set_cell_border(hdr[i], accent)
        set_cell_margins(hdr[i], 105, 105, 105, 105)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        r.font.size = Pt(8.8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                cells[i].width = Cm(widths[i])
            set_cell_border(cells[i])
            set_cell_margins(cells[i], 95, 100, 95, 100)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            r.font.size = Pt(8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, title, text, fill=LIGHT_ORANGE, accent=ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent)
    set_cell_margins(cell, 135, 165, 135, 165)
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    r.font.size = Pt(9.8)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_tight_heading(doc, text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.2)
    r.font.color.rgb = RGBColor.from_string(color)


def create_diagram(path, title, boxes, links=None, colors=None):
    img = Image.new("RGB", (1400, 760), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_big = ImageFont.truetype("arialbd.ttf", 38)
        font = ImageFont.truetype("arial.ttf", 27)
        font_small = ImageFont.truetype("arial.ttf", 22)
    except Exception:
        font_big = font = font_small = ImageFont.load_default()
    draw.rectangle((0, 0, 1400, 760), fill="#F8FAFC")
    draw.text((50, 34), title, fill="#" + GREEN, font=font_big)
    palette = colors or ["#EAF6EE", "#EAF2FB", "#FFF3E5", "#FDECEC"]
    for idx, (label, x, y, w, h) in enumerate(boxes):
        fill = palette[idx % len(palette)]
        outline = ["#" + GREEN, "#" + BLUE, "#" + ORANGE, "#" + RED][idx % 4]
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline=outline, width=4)
        words = label.split()
        lines = []
        line = ""
        for word in words:
            trial = (line + " " + word).strip()
            if len(trial) > 22:
                lines.append(line)
                line = word
            else:
                line = trial
        if line:
            lines.append(line)
        yy = y + h / 2 - len(lines) * 17
        for line in lines:
            tw = draw.textlength(line, font=font)
            draw.text((x + w / 2 - tw / 2, yy), line, fill="#1F2937", font=font)
            yy += 35
    if links:
        for x1, y1, x2, y2 in links:
            draw.line((x1, y1, x2, y2), fill="#64748B", width=5)
            draw.polygon([(x2, y2), (x2 - 18, y2 - 9), (x2 - 18, y2 + 9)], fill="#64748B")
    draw.text((52, 708), "Designed for offline-first rural health support with simple navigation and emergency readiness.", fill="#64748B", font=font_small)
    img.save(path, quality=95)


def make_assets():
    create_diagram(
        ASSETS / "architecture.png",
        "Application Architecture",
        [
            ("User Interface Fragments", 70, 150, 310, 120),
            ("ViewModels and LiveData", 525, 150, 310, 120),
            ("Repositories", 980, 150, 310, 120),
            ("Room Database", 300, 430, 310, 120),
            ("Voice and Emergency Services", 790, 430, 360, 120),
        ],
        links=[(380, 210, 525, 210), (835, 210, 980, 210), (1135, 270, 970, 430), (680, 270, 510, 430), (680, 270, 790, 430)],
    )
    create_diagram(
        ASSETS / "workflow.png",
        "Core User Workflow",
        [
            ("Onboarding Profile", 70, 165, 250, 110),
            ("Home Dashboard", 390, 165, 250, 110),
            ("Medicine and Vital Logging", 710, 165, 285, 110),
            ("ASHA and AI Help", 1065, 165, 250, 110),
            ("Emergency SOS", 565, 450, 280, 110),
        ],
        links=[(320, 220, 390, 220), (640, 220, 710, 220), (995, 220, 1065, 220), (835, 275, 705, 450)],
    )
    create_diagram(
        ASSETS / "database.png",
        "Local Data Model",
        [
            ("MedicalProfile", 90, 160, 250, 95),
            ("Pill", 420, 160, 220, 95),
            ("PillLog", 730, 160, 220, 95),
            ("VitalEntry", 1040, 160, 230, 95),
            ("HealthEvent", 250, 430, 250, 95),
            ("VoiceHistory", 590, 430, 260, 95),
            ("HealthTip", 940, 430, 230, 95),
        ],
        links=[(640, 207, 730, 207)],
    )


def section_page(doc, title, subtitle, bullets, table_rows=None, note=None, diagram=None, accent=GREEN):
    doc.add_heading(title, level=1)
    add_banner(doc, subtitle, fill=LIGHT_GREEN if accent == GREEN else LIGHT_BLUE)
    if diagram:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(diagram), width=Cm(12.8))
    for b in bullets:
        bullet(doc, b)
    if table_rows:
        add_key_value_table(doc, table_rows)
    add_section_traceability(doc, title, compact=bool(diagram))
    if note:
        add_note(doc, note[0], note[1])
    doc.add_page_break()


def add_cover(doc):
    paragraph(doc, "ArogyaSahaya V2", bold=True, color=GREEN, size=30, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    paragraph(doc, "Professional Project Report", bold=True, color=BLUE, size=17, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_banner(doc, "A rural healthcare companion for chronic disease management, medicine adherence, vital tracking, ASHA coordination, voice guidance, and emergency support.", fill=LIGHT_GREEN)
    add_key_value_table(doc, [
        ("Project Type", "Android mobile application"),
        ("Primary Domain", "Digital health, chronic care, emergency assistance"),
        ("Technology Stack", "Kotlin, Android XML UI, Room database, Hilt, LiveData, OkHttp, SpeechRecognizer, TextToSpeech"),
        ("Prepared For", "Academic project documentation and result presentation"),
        ("Prepared On", date.today().strftime("%d %B %Y")),
    ], header="Report Metadata")
    add_note(doc, "Report Scope", "This document explains the purpose, design, implementation, testing, and visible results of ArogyaSahaya V2. Screenshots are included in the Results section as practical evidence of the completed mobile interface.")
    add_matrix_table(doc, ["Report Part", "Coverage"], [
        ("Analysis", "Problem statement, objectives, existing system need, and proposed system scope."),
        ("Design", "Architecture, data model, module decomposition, and interface principles."),
        ("Build", "Android implementation details, voice assistant flow, and emergency safety behavior."),
        ("Validation", "Testing strategy, screenshots, result observations, discussion, and future scope."),
    ], widths=[4.0, 11.2], accent=BLUE)
    paragraph(doc, "The report is formatted with consistent margins, headers, tables, figure captions, and section summaries so it can be submitted directly as a professional project report.", align=WD_ALIGN_PARAGRAPH.CENTER, color="4B5563", size=10.5)
    doc.add_page_break()


def add_front_matter(doc):
    doc.add_heading("Certificate", level=1)
    add_banner(doc, "Project completion statement", "This page may be signed by the guide, evaluator, or department representative after review.")
    for t in [
        "This is to certify that the project titled ArogyaSahaya V2 has been prepared as a complete Android application report describing the design and implementation of a patient assistance system.",
        "The work documented in this report covers problem analysis, requirements, system architecture, database design, feature implementation, testing, screenshots, observations, and future enhancement opportunities.",
        "The application demonstrates key health support functions such as medicine scheduling, vital logging, emergency contact support, voice-based assistance, and ASHA/community health coordination.",
    ]:
        paragraph(doc, t)
    add_key_value_table(doc, [
        ("Project Title", "ArogyaSahaya V2"),
        ("Student / Team", "To be filled by the submitting student or team"),
        ("Guide / Faculty", "To be filled by the institution"),
        ("Department", "Computer Science / Information Technology / Relevant Department"),
        ("Institution", "To be filled by the institution"),
    ])
    add_matrix_table(doc, ["Role", "Name", "Signature", "Date"], [
        ("Student", "", "", ""),
        ("Guide", "", "", ""),
        ("Evaluator", "", "", ""),
    ], widths=[3.2, 5.0, 4.0, 3.0])
    add_matrix_table(doc, ["Evaluation Area", "Expected Evidence"], [
        ("Application purpose", "Clear healthcare problem and user group are documented."),
        ("Technical design", "Android architecture, database entities, and services are explained."),
        ("Implementation", "Medicine, vitals, SOS, voice, profile, and ASHA modules are covered."),
        ("Result proof", "Screenshots are placed in the Results section with observations."),
    ], widths=[4.5, 10.7], accent=BLUE)
    doc.add_page_break()

    doc.add_heading("Declaration and Acknowledgement", level=1)
    add_banner(doc, "Authenticity and gratitude", "A concise front-matter section for formal submission.")
    paragraph(doc, "I declare that this report presents the ArogyaSahaya V2 project in a structured and original form based on the application source files and observed output screens. The report is intended for academic explanation, evaluation, and demonstration.")
    paragraph(doc, "I acknowledge the guidance of faculty members, peers, and users whose practical healthcare needs inspired the design direction. The project focuses on clarity, accessibility, and everyday usefulness for patients who need simple health support.")
    add_key_value_table(doc, [
        ("Main contribution", "A patient-facing Android application that combines medicine reminders, vital tracking, emergency support, and voice-based health guidance."),
        ("Expected users", "Elderly patients, chronic disease patients, caregivers, ASHA workers, and family members."),
        ("Ethical note", "The app provides supportive guidance only. It does not replace qualified medical diagnosis or emergency medical services."),
    ])
    add_matrix_table(doc, ["Principle", "How It Is Reflected"], [
        ("Clarity", "Screens use direct labels, recognizable icons, and short patient-facing wording."),
        ("Safety", "SOS actions are visually dominant and AI advice avoids diagnosis."),
        ("Privacy", "Important health records are stored locally through Room database entities."),
        ("Accessibility", "Voice input and spoken output support users who may struggle with typing."),
        ("Continuity", "Logs and history screens preserve medicines, vitals, and voice interactions."),
    ], widths=[4.0, 11.2], accent=BLUE)
    add_note(doc, "Submission Note", "Names, registration numbers, guide details, and institution information can be added in the reserved fields without disturbing the report layout.")
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_banner(doc, "Project overview in brief", "The application addresses routine care gaps through accessible mobile features.")
    for t in [
        "ArogyaSahaya V2 is an Android healthcare companion designed to help patients manage chronic conditions such as diabetes and blood pressure through simple daily workflows. The application brings together medicine reminders, vital record keeping, emergency calling, ASHA/community event awareness, medical profile storage, and an AI-enabled voice help screen.",
        "The system uses a local Room database to preserve important patient data on the device, while repositories and ViewModels provide a clean flow between storage and user interface screens. The app also integrates Android speech recognition and text-to-speech to make health help more approachable for users who may prefer voice interaction.",
        "The result is a practical mobile solution that prioritizes visibility, large touch targets, simple language, and fast access to emergency actions. The report documents the problem background, requirements, system design, implementation details, testing strategy, screenshots, limitations, and future improvements.",
    ]:
        paragraph(doc, t)
    add_matrix_table(doc, ["Keyword", "Meaning in Project"], [
        ("Arogya", "Health and wellness support"),
        ("Medicine adherence", "Tracking timely medicine intake and remaining stock"),
        ("Vitals", "Blood pressure, heart rate, glucose, oxygen, temperature, and related records"),
        ("SOS", "Emergency call, contact call, and SMS/location support"),
        ("Voice help", "Health question interface with online AI and offline fallback responses"),
    ], widths=[4.2, 11.0])
    add_note(doc, "Abstract Outcome", "The prototype demonstrates a complete patient workflow: view dashboard, manage medicine, log vitals, ask a health question, and request emergency help.")
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    add_banner(doc, "Report navigation", "Static contents for quick reading and evaluation.")
    toc = [
        "1. Introduction", "2. Problem Statement", "3. Objectives", "4. Existing System and Need", "5. Proposed System",
        "6. Requirement Specification", "7. Technology Stack", "8. System Architecture", "9. Database Design",
        "10. Module Design", "11. User Interface Design", "12. Implementation Details", "13. Voice Assistant",
        "14. Emergency and Safety Design", "15. Testing Strategy", "16. Results and Screenshots", "17. Discussion",
        "18. Future Scope", "19. Conclusion", "20. References"
    ]
    rows = [(item, "Included") for item in toc]
    add_matrix_table(doc, ["Section", "Status"], rows, widths=[12.2, 3.0])
    add_note(doc, "Layout Note", "The report uses page-level sectioning so each major topic is easy to read and no page is left as an empty spacer.")
    doc.add_page_break()


def add_results_page(doc, idx, title, image_path, observations):
    doc.add_heading(f"Result {idx}: {title}", level=1)
    add_banner(doc, "Observed application output", "Screenshot evidence from the completed Android interface.", fill=LIGHT_BLUE)
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), height=Cm(13.7))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(f"Figure {idx}: {title}")
        cr.bold = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor.from_string("4B5563")
    else:
        add_note(doc, "Missing Screenshot", f"The source image was not found at {image_path}.", fill="FDECEC", accent=RED)
    add_matrix_table(doc, ["Evaluation Point", "Result"], observations, widths=[5.0, 10.2], accent=BLUE)
    add_note(doc, "Result Summary", "The screen follows the application goal of keeping health actions visible, direct, and simple for daily patient use.", fill=LIGHT_GREEN, accent=GREEN)
    doc.add_page_break()


def build_doc():
    make_assets()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_front_matter(doc)

    section_page(doc, "1. Introduction", "Healthcare support problem addressed by the application", [
        "Chronic disease patients often need repeated reminders for medicines, periodic vital checks, and fast access to emergency support.",
        "ArogyaSahaya V2 brings these daily needs into one mobile application with a simple navigation model.",
        "The app is especially suitable for elderly and rural users because the interface uses clear labels, familiar health symbols, and voice interaction.",
        "The project demonstrates how local data storage and Android system services can create a dependable patient support tool.",
    ], [
        ("Target condition support", "Diabetes, blood pressure, routine medication, and emergency assistance workflows."),
        ("Primary interaction style", "Touch-first screens supported by microphone input and text-to-speech output."),
        ("Core design direction", "Simple language, card-based information display, and prominent action buttons."),
    ], ("Key Point", "The application is a care-support tool. It helps users remember, record, and respond, while serious health decisions remain with doctors and emergency services."))

    section_page(doc, "2. Problem Statement", "Care gaps that motivated ArogyaSahaya V2", [
        "Many patients forget doses or lose track of medicine stock, especially when more than one medicine is prescribed.",
        "Vital readings are often written on paper or remembered verbally, which makes trends difficult to review.",
        "Emergency help must be reachable in very few taps because delay can increase risk during dizziness, chest pain, low sugar, or sudden weakness.",
        "Patients may need basic health guidance in simple language before deciding whether to visit a doctor.",
    ], [
        ("Medicine adherence gap", "Missed doses, unclear timing, and poor refill awareness."),
        ("Health record gap", "No convenient consolidated log for blood pressure, glucose, heart rate, and oxygen."),
        ("Emergency gap", "Family contact, ambulance number, and safety tips may not be available quickly."),
        ("Communication gap", "Users may prefer voice questions instead of typing long medical concerns."),
    ], ("Problem Focus", "The project concentrates on everyday patient assistance rather than hospital administration or doctor-side diagnosis."))

    section_page(doc, "3. Objectives", "Functional and quality goals", [
        "Provide a home dashboard that summarizes adherence, latest vitals, medicines, ASHA connect, medical profile, and emergency access.",
        "Allow users to add medicines with dose timing, instruction, stock, and refill-related information.",
        "Allow users to record vital readings and review recent history and trends.",
        "Offer emergency options such as SOS hold action, emergency contact calling, 112 call action, and SMS support.",
        "Provide a voice assistant that answers common health questions using AI when configured and offline fallback responses otherwise.",
    ], [
        ("Usability objective", "Keep primary actions visible and reachable through bottom navigation."),
        ("Reliability objective", "Store profile, medicine, vital, event, and voice history records locally with Room."),
        ("Safety objective", "Use clear emergency warnings and provide direct contact options."),
        ("Scalability objective", "Keep modules separated so new screens and data entities can be added later."),
    ], ("Success Criteria", "The app is successful when a user can open it, understand current health tasks, record information, and reach help without training."))

    section_page(doc, "4. Existing System and Need", "Comparison with manual and fragmented workflows", [
        "Manual medicine charts can be lost, forgotten, or updated inconsistently.",
        "Generic reminder apps do not understand health-specific fields such as dose time, stock, refill alert, and adherence calculation.",
        "Paper-based vital logs do not automatically support latest-reading display or trend views.",
        "Emergency contact information may be buried inside the phone contacts app, slowing the response workflow.",
    ], [
        ("Manual record keeping", "Low cost but hard to search, trend, or share."),
        ("General alarm apps", "Useful for time alerts but weak for medicine context and adherence history."),
        ("Separate health apps", "Often complex for elderly users and may not combine SOS and voice help."),
        ("ArogyaSahaya V2", "Unifies patient support tasks in a simple Android interface."),
    ], ("Need Statement", "A unified, local-first, patient-facing system improves continuity between reminders, records, advice, and emergency readiness."))

    section_page(doc, "5. Proposed System", "Integrated Android healthcare companion", [
        "The proposed system uses a single Android application with bottom navigation across Home, Medicines, Vitals, AI Help, and SOS.",
        "Room entities represent medical profile, medicines, dose logs, vital entries, health events, voice history, and health tips.",
        "The UI uses fragments for modular screens and ViewModels for lifecycle-aware data presentation.",
        "Emergency features rely on Android permissions such as CALL_PHONE, SEND_SMS, VIBRATE, and location-ready communication workflows.",
    ], [
        ("Home", "Daily overview, adherence status, latest vitals, and shortcuts."),
        ("Medicines", "Active medicine list, dose timing, stock, and taken logging."),
        ("Vitals", "Latest reading card, trend chart, and chronological history."),
        ("AI Help", "Microphone-based health question screen with fallback guidance."),
        ("SOS", "Large emergency action, contact details, emergency calls, SMS, and tips."),
    ], ("Design Choice", "The system keeps patient-facing workflows separate but connected through a shared local database and common navigation."))

    section_page(doc, "6. Requirement Specification", "Functional and non-functional requirements", [
        "The application must save and update a single medical profile for the user.",
        "The application must create, update, deactivate, and log medicines with dose timing.",
        "The application must save vital entries and display the latest reading and recent trend data.",
        "The application must preserve voice question history and prune old records to avoid unnecessary data growth.",
        "The application should remain understandable even when online AI configuration is not available.",
    ], [
        ("Functional requirement", "Add medicine, log dose, add vitals, call emergency contact, ask voice question."),
        ("Performance requirement", "Local lists and latest values should load quickly from Room and LiveData."),
        ("Usability requirement", "Important screens must use simple labels and large visual affordances."),
        ("Safety requirement", "AI responses must include doctor/emergency escalation language for serious conditions."),
        ("Maintainability requirement", "Modules should stay separated across UI, repository, database, and service layers."),
    ], ("Requirement Priority", "Emergency access, medicine adherence, and vital logging are the highest priority because they directly affect patient safety."))

    section_page(doc, "7. Technology Stack", "Tools, frameworks, and Android services", [
        "Kotlin is used for Android application logic, fragments, ViewModels, repositories, services, and receivers.",
        "Android XML layouts define the visual interface, cards, buttons, dialogs, and navigation screens.",
        "Room database provides type-safe local persistence for structured health data.",
        "Hilt dependency injection connects database and repository objects without manual wiring in every screen.",
        "OkHttp and JSON are used by the voice assistant service to call the Gemini API when a key is configured.",
    ], [
        ("Programming language", "Kotlin"),
        ("UI layer", "Android XML layouts, fragments, bottom navigation"),
        ("Persistence", "Room database with DAO interfaces and LiveData queries"),
        ("Dependency injection", "Hilt modules"),
        ("Voice", "SpeechRecognizer and TextToSpeech"),
        ("Networking", "OkHttp with JSON request and response parsing"),
    ], ("Stack Fit", "The selected stack is appropriate because Android provides native access to alarms, calls, SMS, speech input, text-to-speech, and local storage."))

    section_page(doc, "8. System Architecture", "Layered flow from screen to storage and services", [
        "UI fragments collect user actions and display current data.",
        "ViewModels expose lifecycle-aware state and isolate UI screens from database details.",
        "Repositories provide a single access point for DAO operations and service history saving.",
        "Room DAOs execute structured local queries for medicines, vitals, events, profile, and voice records.",
        "Android services and receivers support voice response and medicine reminder behavior.",
    ], [
        ("Presentation layer", "HomeFragment, PillFragment, VitalsFragment, EmergencyFragment, VoiceFragment, ProfileFragment, AshaFragment."),
        ("State layer", "ViewModels hold screen state and call repositories."),
        ("Data layer", "Repositories, DAOs, Room database, and entity models."),
        ("System layer", "AlarmReceiver, VoiceAssistantService, Android permissions, speech APIs."),
    ], ("Architecture Benefit", "A layered structure makes the project easier to test, explain, and extend."), ASSETS / "architecture.png")

    section_page(doc, "9. Database Design", "Room entities and DAO responsibilities", [
        "The database version contains seven entities: MedicalProfile, Pill, PillLog, VitalEntry, HealthEvent, VoiceHistory, and HealthTip.",
        "The profile table stores demographics, chronic conditions, allergies, emergency contact, doctor contact, and preferred language.",
        "The medicine tables separate configured medicines from dose log entries, which supports adherence calculation.",
        "The vitals table records multiple health measurements and can return latest, all, last seven, or date-filtered records.",
        "Voice and health-tip tables preserve recent interactions while pruning older entries.",
    ], [
        ("MedicalProfile", "Single profile row with patient identity and emergency information."),
        ("Pill and PillLog", "Medication schedule, stock, reminders, and taken history."),
        ("VitalEntry", "BP, heart rate, glucose, oxygen, temperature, weight, notes, and mood."),
        ("HealthEvent", "ASHA visits, health camps, location, organizer, and completion status."),
        ("VoiceHistory", "Recent voice queries and responses for continuity."),
    ], ("Database Strength", "The schema is compact but covers the most important patient support data without depending on internet availability."), ASSETS / "database.png")

    section_page(doc, "10. Module Design", "Screen-level decomposition", [
        "The Home module summarizes important information and provides shortcuts to major health workflows.",
        "The Medicines module handles active medicine display, add/edit dialogs, dose confirmation, and stock visibility.",
        "The Vitals module manages numeric health inputs, latest reading status, trend presentation, and history.",
        "The ASHA module can list upcoming health camps or worker visits through the health events data model.",
        "The Profile module stores patient and contact details that support emergency and personalization features.",
    ], [
        ("Home module", "Patient greeting, condition status, adherence progress, latest vitals, and feature cards."),
        ("Medicine module", "Medicine list, dose labels, stock count, and dose taken action."),
        ("Vitals module", "Latest BP/glucose/heart/oxygen view and history records."),
        ("Voice module", "Question input, AI response, offline fallback, TTS output, and history."),
        ("Emergency module", "SOS button, contact call, 112 call, SMS, and emergency tips."),
    ], ("Module Principle", "Each module does one visible job, but all modules share the same patient context."))

    section_page(doc, "11. User Interface Design", "Accessibility, consistency, and patient focus", [
        "The interface uses large cards, strong color coding, and bottom navigation to reduce confusion.",
        "Health domains are mapped to intuitive colors: green for care and adherence, blue for vitals, red for emergency, and orange for alerts.",
        "Cards contain concise status text, icon cues, and action buttons so users do not need to read long instructions.",
        "Floating action buttons support quick additions such as logging vitals or adding medicine.",
        "The SOS screen intentionally uses a large red action target because emergency use should be unmistakable.",
    ], [
        ("Navigation", "Five-tab bottom bar: Home, Medicines, Vitals, AI Help, SOS."),
        ("Typography", "Large screen titles, readable body text, and short labels."),
        ("Touch design", "Large buttons for patient actions and emergency calls."),
        ("Feedback", "Toast-style confirmation, adherence progress, latest readings, and screen status."),
    ], ("UI Standard", "The design is intentionally simple, not decorative. It prioritizes comprehension and speed."))

    section_page(doc, "12. Implementation Details", "Important source-level behaviors", [
        "AndroidManifest declares permissions for exact alarms, boot completed handling, notifications, vibration, wake lock, calling, SMS, foreground service, microphone, internet, network state, and biometric support.",
        "AlarmReceiver is registered for pill alarms and boot completed events so reminders can continue after device restart.",
        "VoiceAssistantService is an Android service that initializes TextToSpeech, starts speech recognition, processes results, and stores response history.",
        "DAO methods provide LiveData for automatic UI updates and suspend functions for insert, update, delete, and latest-record operations.",
        "The Gemini integration is prepared through an API key placeholder while offline fallback responses keep the app usable without configuration.",
    ], [
        ("Manifest evidence", "CALL_PHONE, SEND_SMS, RECORD_AUDIO, INTERNET, POST_NOTIFICATIONS, and alarm permissions are declared."),
        ("Local data", "Room database version 1 with exportSchema disabled for compact project setup."),
        ("Voice fallback", "Common topics include blood pressure, diabetes, medicine, pain, sleep, diet, exercise, and emergency."),
        ("Reminder support", "Medicine model includes morning, afternoon, and night reminder fields."),
    ], ("Implementation Note", "The report describes the implemented behavior visible in the source and screenshots, while API keys and institution-specific details remain configurable."))

    section_page(doc, "13. Voice Assistant Design", "Arogya Mitra health question workflow", [
        "The AI Help screen offers a microphone-centered interface for users who may find typing difficult.",
        "SpeechRecognizer captures patient questions in Indian English language settings.",
        "The service builds a patient-aware prompt using the patient name and known conditions.",
        "Responses are kept concise and supportive, with explicit advice to consult doctors for serious concerns.",
        "If the online Gemini call fails or the API key is not configured, the service returns safe offline guidance for common questions.",
    ], [
        ("Input", "Microphone tap starts listening and returns recognized text."),
        ("Processing", "Gemini API request is attempted with safety-oriented system instructions."),
        ("Fallback", "Rule-based response covers BP, sugar, medicine, pain, sleep, diet, exercise, and emergency."),
        ("Output", "Response is displayed, spoken using TTS, and saved to voice history."),
    ], ("Safety Rule", "The assistant is framed as a helper, not a diagnosing doctor. Emergency symptoms route users toward 112 or medical care."))

    section_page(doc, "14. Emergency and Safety Design", "Fast help access and careful guidance", [
        "The SOS screen is visually distinct with red color, a large center button, and short emergency instructions.",
        "The app supports emergency contact display, doctor contact placeholders, direct 112 calling, contact calling, and SOS SMS with location intent.",
        "Emergency tips remind users to stay calm, lie down if dizzy, avoid extra medicine without advice, and keep the profile card visible.",
        "The manifest includes call and SMS permissions, which are necessary for emergency communication features.",
        "The design reduces navigation effort by placing SOS permanently in the bottom bar.",
    ], [
        ("Primary action", "Large SOS hold button."),
        ("Secondary action", "Call 112 and call saved contact."),
        ("Message action", "Send SOS SMS with location."),
        ("Guidance", "Emergency tips shown directly on the screen."),
    ], ("Safety Outcome", "The emergency module turns stored profile/contact data into immediate action options."))

    section_page(doc, "15. Testing Strategy", "Verification approach for a patient-facing app", [
        "Functional testing checks that medicine entries, vital entries, emergency actions, profile details, and voice queries execute correctly.",
        "Database testing checks insert, update, delete, latest item, active item, range query, and pruning behavior.",
        "UI testing checks readability, button placement, bottom navigation, screen state, and dialog behavior on mobile dimensions.",
        "Permission testing checks app behavior when microphone, calls, SMS, notifications, or exact alarms are denied.",
        "Safety testing checks that emergency and AI fallback messages do not provide dangerous or overconfident medical diagnosis.",
    ], [
        ("Unit-level checks", "DAO queries, repository calls, fallback response mapping."),
        ("Integration checks", "Fragment to ViewModel to repository to Room flow."),
        ("Manual UI checks", "Add medicine, mark dose, add vitals, open SOS, ask AI question."),
        ("Edge cases", "No profile, no active medicines, empty vitals, no internet, missing API key."),
        ("Acceptance checks", "Screenshots confirm main flows are present and readable."),
    ], ("Testing Priority", "For this project, the highest-risk areas are emergency workflows, reminder persistence, and health advice wording."))

    # Result pages with screenshots.
    result_rows = [
        [
            ("Latest readings", "Displays BP 120/80, glucose 115.0, heart rate 75, and oxygen 99.0."),
            ("Trend evidence", "Shows a 7-day trend area and history card for the logged reading."),
            ("Usability", "The add vital action remains visible through a floating button."),
        ],
        [
            ("Medicine record", "Shows one active medicine named Dollo with tablet type and after-food instruction."),
            ("Adherence action", "Provides a clear morning dose confirmation button."),
            ("Stock awareness", "Displays stock count so refill planning is visible."),
        ],
        [
            ("Emergency readiness", "Large SOS button, call 112, call contact, and SOS SMS options are visible."),
            ("Contact support", "Emergency contact name and phone are displayed for direct action."),
            ("Safety guidance", "Emergency tips are presented below the action controls."),
        ],
        [
            ("Dashboard overview", "Greets the patient, shows condition management, adherence, vitals, and feature cards."),
            ("Navigation", "Home, medicines, vitals, AI help, and SOS are reachable from the bottom bar."),
            ("Patient focus", "Large shortcut cards simplify access to common health actions."),
        ],
        [
            ("Voice interface", "Arogya Mitra provides a microphone-centered health question screen."),
            ("Quick questions", "Common question chips reduce typing effort."),
            ("Offline transparency", "The screen clearly indicates offline mode when API configuration is pending."),
        ],
    ]
    for idx, ((title, path), rows) in enumerate(zip(SCREENSHOTS, result_rows), start=1):
        add_results_page(doc, idx, title, path, rows)

    section_page(doc, "17. Discussion", "What the results show", [
        "The screenshots confirm that the major workflows are implemented and reachable from a consistent navigation structure.",
        "The home screen works as a command center by summarizing adherence, vitals, medicine, ASHA connect, profile, and emergency access.",
        "The medicine and vitals screens show that the app supports both planned health routines and measured health records.",
        "The SOS screen is appropriately dominant and action-oriented, which is important for a safety-critical feature.",
        "The AI Help screen shows a practical path toward voice-first health guidance while still supporting offline behavior.",
    ], [
        ("Strength", "Integrated daily care and emergency support in one app."),
        ("Strength", "Local storage makes the core app useful without permanent internet dependency."),
        ("Limitation", "Online AI needs a configured Gemini API key for full responses."),
        ("Limitation", "Clinical validation and real patient usability testing should be completed before deployment."),
    ], ("Interpretation", "The project has reached a strong prototype stage with clear feature coverage and demonstrable UI output."))

    section_page(doc, "18. Future Scope", "Enhancements for real-world deployment", [
        "Add secure cloud backup so a patient can restore data after changing phones.",
        "Add caregiver and doctor sharing modes with consent-based access.",
        "Add multilingual UI and voice response packs for Kannada, Hindi, and other regional languages.",
        "Add charts for glucose, blood pressure, and adherence over longer periods.",
        "Integrate location services more deeply for emergency SMS and nearest facility suggestions.",
        "Add biometric-protected profile export for clinic visits.",
    ], [
        ("Near-term enhancement", "Polish reminder scheduling, notification channels, and refill alerts."),
        ("Medium-term enhancement", "Add PDF health report export for doctors and ASHA workers."),
        ("Long-term enhancement", "Deploy secure backend sync, caregiver portal, and analytics dashboard."),
        ("Validation need", "Conduct usability testing with elderly and rural users."),
    ], ("Roadmap", "Future work should preserve the app's simplicity while adding reliability, privacy, and sharing features."))

    section_page(doc, "19. Conclusion", "Final project summary", [
        "ArogyaSahaya V2 successfully demonstrates a mobile healthcare companion centered on everyday patient support.",
        "The project combines medicine management, vital logging, ASHA/community event awareness, medical profile storage, AI-based voice support, and emergency actions.",
        "The architecture is understandable and maintainable because source responsibilities are separated into UI fragments, ViewModels, repositories, DAOs, Room entities, services, and receivers.",
        "The result screenshots show that the application interface is complete enough for academic demonstration and further refinement.",
        "With additional validation, security hardening, multilingual polish, and deployment testing, the project can become a practical assistive tool for chronic care users.",
    ], [
        ("Project value", "Improves continuity between reminders, records, advice, and emergency response."),
        ("Academic value", "Demonstrates Android development with local database, services, permissions, and UI modules."),
        ("Social value", "Targets patients who need simple, accessible, and reliable health assistance."),
    ], ("Final Note", "The project should be presented as supportive healthcare technology, not as a replacement for medical professionals."))

    doc.add_heading("20. References", level=1)
    add_banner(doc, "Sources and technical references", "Project source files and platform documentation categories used for this report.")
    refs = [
        ("Android Developers", "Fragments, services, permissions, SpeechRecognizer, TextToSpeech, notifications, alarms, and navigation guidance."),
        ("Android Jetpack", "Room database, LiveData, ViewModel, and lifecycle-aware architecture patterns."),
        ("Kotlin", "Application logic, data classes, coroutines, and Android service implementation."),
        ("OkHttp and JSON", "Network request construction for Gemini API integration."),
        ("Project source files", "ArogyaSahayaV2 AndroidManifest, model classes, Room database DAOs, UI fragments, and VoiceAssistantService."),
        ("Result screenshots", "Mobile output screens supplied for Home, Medicines, Vitals, AI Help, and Emergency SOS."),
    ]
    add_matrix_table(doc, ["Reference Area", "Use in Report"], refs, widths=[4.5, 10.7])
    add_tight_heading(doc, "Glossary of project terms", color=BLUE)
    add_matrix_table(doc, ["Term", "Meaning"], [
        ("Adherence", "The percentage or record of medicines taken as scheduled."),
        ("ASHA", "Accredited Social Health Activist, a community health worker role in India."),
        ("DAO", "Data Access Object used by Room to define database queries."),
        ("Fallback response", "A safe offline answer shown when the online AI service is unavailable."),
        ("LiveData", "Lifecycle-aware observable data used to update Android screens."),
        ("SOS", "Emergency action path for urgent call, contact, and SMS workflows."),
    ], widths=[4.0, 11.2], accent=BLUE)
    add_note(doc, "Professional Use", "Before external deployment, add privacy policy, consent workflow, API key security, clinical review, and field testing documentation.")
    paragraph(doc, "End of Report", bold=True, color=GREEN, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_header_footer(doc)
    DOCX.parent.mkdir(exist_ok=True)
    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    path = build_doc()
    print(path)
